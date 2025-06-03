# src/attachment_processor.py
"""附件处理模块 - 专门处理简历附件解析 - 改进版，支持多种Excel格式"""

import os
import io
import json
import logging
import re
from typing import List, Dict, Optional
from datetime import datetime, timedelta
import base64

try:
    import docx

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not installed. Word document processing disabled.")

try:
    import PyPDF2

    PDF_AVAILABLE = True
    PDF_LIBRARY = "PyPDF2"
except ImportError:
    try:
        import pdfplumber

        PDF_AVAILABLE = True
        PDF_LIBRARY = "pdfplumber"
    except ImportError:
        PDF_AVAILABLE = False
        logging.warning(
            "Neither PyPDF2 nor pdfplumber installed. PDF processing disabled."
        )

try:
    import openpyxl

    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logging.warning("openpyxl not installed. Modern Excel (.xlsx) processing disabled.")

try:
    import xlrd

    XLRD_AVAILABLE = True
except ImportError:
    XLRD_AVAILABLE = False
    logging.warning("xlrd not installed. Legacy Excel (.xls) processing disabled.")

try:
    import pandas as pd

    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    logging.warning("pandas not installed. Alternative Excel processing disabled.")

import httpx
from openai import AsyncOpenAI
from pydantic import BaseModel, Field, field_validator

logger = logging.getLogger(__name__)


class ResumeData(BaseModel):
    """简历数据模型 - 完善的类型转换版本"""

    name: str
    email: Optional[str] = None
    phone: Optional[str] = None
    gender: Optional[str] = None
    age: Optional[str] = None
    nationality: Optional[str] = None
    nearest_station: Optional[str] = None
    education: Optional[str] = None
    arrival_year_japan: Optional[str] = None
    certifications: List[str] = Field(default_factory=list)
    skills: List[str] = Field(default_factory=list)
    technical_keywords: List[str] = Field(default_factory=list)
    experience: str = ""
    work_scope: Optional[str] = None
    work_experience: Optional[str] = None
    japanese_level: Optional[str] = None
    english_level: Optional[str] = None
    availability: Optional[str] = None
    preferred_work_style: List[str] = Field(default_factory=list)
    preferred_locations: List[str] = Field(default_factory=list)
    desired_rate_min: Optional[int] = None
    desired_rate_max: Optional[int] = None
    overtime_available: Optional[bool] = False
    business_trip_available: Optional[bool] = False
    self_promotion: Optional[str] = None
    remarks: Optional[str] = None
    recommendation: Optional[str] = None
    source_filename: Optional[str] = None

    @field_validator("name", mode="before")
    @classmethod
    def validate_name(cls, v):
        """姓名验证器"""
        if not v or v is None:
            return "名前不明"
        return str(v)

    @field_validator("experience", mode="before")
    @classmethod
    def validate_experience(cls, v):
        """经验验证器"""
        if not v or v is None:
            return "不明"
        return str(v)

    @field_validator("age", mode="before")
    @classmethod
    def validate_age(cls, v):
        """年龄验证器 - 支持多种类型转换"""
        if v is None:
            return None

        # 如果是数字，直接转换为字符串
        if isinstance(v, (int, float)):
            return str(int(v))

        # 如果是字符串，提取数字部分
        if isinstance(v, str):
            # 提取数字
            numbers = re.findall(r"\d+", v)
            if numbers:
                return numbers[0]
            return v

        return str(v)

    @field_validator("phone", mode="before")
    @classmethod
    def validate_phone(cls, v):
        """电话号码验证器"""
        if v is None:
            return None
        if isinstance(v, (int, float)):
            return str(int(v))
        return str(v)

    @field_validator("arrival_year_japan", mode="before")
    @classmethod
    def validate_arrival_year_japan(cls, v):
        """来日年度验证器 - 处理Excel日期序列号"""
        if v is None:
            return None

        # 如果是Excel日期序列号（浮点数），转换为年份
        if isinstance(v, float):
            try:
                # Excel日期序列号转换（1900年1月1日为基准）
                # 42465.0 大约对应 2016年
                base_date = datetime(1900, 1, 1)
                # Excel有一个闰年bug，需要减去2天
                actual_date = base_date + timedelta(days=int(v) - 2)
                return str(actual_date.year)
            except:
                # 如果转换失败，尝试直接当作年份
                year = int(v)
                if 1900 <= year <= 2100:
                    return str(year)
                elif year > 40000:  # 可能是Excel序列号
                    return str(2000 + (year - 40000) // 365)  # 简单估算
                return str(year)

        # 如果是整数，直接转换
        if isinstance(v, int):
            if 1900 <= v <= 2100:
                return str(v)
            elif v > 40000:  # Excel序列号
                return str(2000 + (v - 40000) // 365)
            return str(v)

        # 如果是字符串，提取年份
        if isinstance(v, str):
            numbers = re.findall(r"\d{4}", v)  # 查找4位数年份
            if numbers:
                return numbers[0]
            # 查找2位数年份
            numbers = re.findall(r"\d{2}", v)
            if numbers:
                year = int(numbers[0])
                if year < 50:
                    return str(2000 + year)
                else:
                    return str(1900 + year)
            return v

        return str(v)

    @field_validator("gender", mode="before")
    @classmethod
    def validate_gender(cls, v):
        """性别验证器"""
        if v is None:
            return None

        v_str = str(v).lower()

        if any(word in v_str for word in ["男", "male", "m"]):
            return "男性"
        elif any(word in v_str for word in ["女", "female", "f"]):
            return "女性"
        else:
            return "回答しない"

    @field_validator("japanese_level", "english_level", mode="before")
    @classmethod
    def validate_language_level(cls, v):
        """语言水平验证器"""
        if v is None:
            return None

        v_str = str(v).lower()

        # 语言水平映射规则
        mappings = {
            "n1": "ネイティブレベル",
            "n2": "ビジネスレベル",
            "n3": "日常会話レベル",
            "n4": "日常会話レベル",
            "n5": "日常会話レベル",
            "1級": "ネイティブレベル",
            "2級": "ビジネスレベル",
            "ネイティブ": "ネイティブレベル",
            "native": "ネイティブレベル",
            "ビジネス": "ビジネスレベル",
            "business": "ビジネスレベル",
            "日常会話": "日常会話レベル",
            "conversational": "日常会話レベル",
            "不問": "不問",
            "なし": "不問",
            "none": "不問",
        }

        for key, mapped_value in mappings.items():
            if key in v_str:
                return mapped_value

        return "日常会話レベル"  # 默认值

    @field_validator(
        "preferred_work_style",
        "preferred_locations",
        "certifications",
        "skills",
        "technical_keywords",
        mode="before",
    )
    @classmethod
    def validate_list_fields(cls, v):
        """列表字段验证器 - 处理None值和各种输入类型"""
        if v is None:
            return []
        if isinstance(v, list):
            return [str(item) for item in v if item is not None]
        if isinstance(v, str):
            if v.strip() == "":
                return []
            # 尝试按逗号分割
            items = [item.strip() for item in v.split(",") if item.strip()]
            return items
        # 如果是其他类型，转换为字符串并放入列表
        return [str(v)]

    @field_validator("desired_rate_min", "desired_rate_max", mode="before")
    @classmethod
    def validate_rate(cls, v):
        """单价验证器"""
        if v is None:
            return None
        if isinstance(v, (int, float)):
            return int(v)
        if isinstance(v, str):
            # 从字符串中提取数字
            numbers = re.findall(r"\d+", v)
            if numbers:
                return int(numbers[0])
            return None
        return None

    @field_validator("overtime_available", "business_trip_available", mode="before")
    @classmethod
    def validate_boolean(cls, v):
        """布尔值验证器"""
        if v is None:
            return False
        if isinstance(v, bool):
            return v
        if isinstance(v, str):
            return v.lower() in ("true", "yes", "可能", "可", "ok", "対応可能", "はい")
        return False

    @field_validator(
        "nationality",
        "nearest_station",
        "education",
        "work_scope",
        "work_experience",
        "availability",
        "self_promotion",
        "remarks",
        "recommendation",
        "source_filename",
        mode="before",
    )
    @classmethod
    def validate_optional_string_fields(cls, v):
        """可选字符串字段验证器"""
        if v is None:
            return None
        if isinstance(v, (int, float)):
            return str(v)
        return str(v)

    @field_validator("email", mode="before")
    @classmethod
    def validate_email(cls, v):
        """邮箱验证器"""
        if v is None:
            return None

        email_str = str(v)
        # 简单的邮箱格式检查
        if "@" in email_str and "." in email_str:
            return email_str

        # 如果不是有效邮箱格式，返回None
        return None


class AttachmentProcessor:
    """附件处理器 - 改进版"""

    def __init__(self, ai_config: Dict):
        self.ai_config = ai_config
        self.ai_client = None

        provider_name = ai_config.get("provider_name")
        api_key = ai_config.get("api_key")

        if provider_name == "openai":
            if api_key:
                self.ai_client = AsyncOpenAI(api_key=api_key)
        elif provider_name in ["deepseek", "custom"]:
            api_base_url = ai_config.get("api_base_url")
            timeout = ai_config.get("timeout", 120.0)
            if api_key and api_base_url:
                self.ai_client = httpx.AsyncClient(
                    base_url=api_base_url,
                    headers={
                        "Authorization": f"Bearer {api_key}",
                        "Content-Type": "application/json",
                    },
                    timeout=timeout,
                )
                logger.info(
                    f"AttachmentProcessor {provider_name.title()} client initialized"
                )

    def detect_file_type(self, file_content: bytes, filename: str) -> str:
        """检测文件的真实类型"""
        if not file_content:
            return "unknown"

        # 检查文件头部魔数
        header = file_content[:8]

        # Excel (.xlsx) - ZIP格式
        if header.startswith(b"PK\x03\x04"):
            return "xlsx"

        # Excel (.xls) - OLE格式
        if header.startswith(b"\xd0\xcf\x11\xe0"):
            return "xls"

        # PDF
        if header.startswith(b"%PDF"):
            return "pdf"

        # Word (.docx) - ZIP格式，需要进一步检查
        if header.startswith(b"PK\x03\x04"):
            try:
                # 尝试作为ZIP文件读取，检查是否包含word相关文件
                import zipfile

                with zipfile.ZipFile(io.BytesIO(file_content)) as zip_file:
                    file_list = zip_file.namelist()
                    if any("word/" in f for f in file_list):
                        return "docx"
                    elif any("xl/" in f for f in file_list):
                        return "xlsx"
            except:
                pass

        # 检查是否为HTML（一些.xls文件实际上是HTML表格）
        try:
            text_content = file_content.decode("utf-8", errors="ignore")[:1000]
            if "<html" in text_content.lower() or "<table" in text_content.lower():
                return "html_table"
        except:
            pass

        # 检查是否为纯文本
        try:
            file_content.decode("utf-8")
            return "text"
        except:
            pass

        return "unknown"

    def extract_text_from_docx(self, file_content: bytes) -> str:
        """从Word文档提取文本"""
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available for Word document processing")
            return ""

        try:
            doc = docx.Document(io.BytesIO(file_content))
            full_text = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())

            # 处理表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())

            text = "\n".join(full_text)
            logger.info(f"从Word文档提取了 {len(text)} 字符的文本")
            return text

        except Exception as e:
            logger.error(f"Word文档解析失败: {e}")
            return ""

    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """从PDF文件提取文本"""
        if not PDF_AVAILABLE:
            logger.error("PDF processing library not available")
            return ""

        try:
            if PDF_LIBRARY == "pdfplumber":
                import pdfplumber

                full_text = []
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            full_text.append(text)
                text = "\n".join(full_text)
                logger.info(f"从PDF文档提取了 {len(text)} 字符的文本")
                return text
            else:  # PyPDF2
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                full_text = []
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        full_text.append(text)
                text = "\n".join(full_text)
                logger.info(f"从PDF文档提取了 {len(text)} 字符的文本")
                return text

        except Exception as e:
            logger.error(f"PDF文档解析失败: {e}")
            return ""

    def extract_text_from_html_table(self, file_content: bytes) -> str:
        """从HTML表格提取文本（一些.xls文件实际上是HTML）"""
        try:
            from bs4 import BeautifulSoup

            html_content = file_content.decode("utf-8", errors="ignore")
            soup = BeautifulSoup(html_content, "html.parser")

            # 提取所有表格内容
            tables = soup.find_all("table")
            full_text = []

            for table in tables:
                rows = table.find_all("tr")
                for row in rows:
                    cells = row.find_all(["td", "th"])
                    row_text = []
                    for cell in cells:
                        cell_text = cell.get_text(strip=True)
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        full_text.append(" | ".join(row_text))

            # 如果没有表格，提取所有文本
            if not full_text:
                text = soup.get_text(separator="\n", strip=True)
                full_text = [line for line in text.split("\n") if line.strip()]

            text = "\n".join(full_text)
            logger.info(f"从HTML表格提取了 {len(text)} 字符的文本")
            return text

        except Exception as e:
            logger.error(f"HTML表格解析失败: {e}")
            # 如果BeautifulSoup不可用，尝试简单的文本提取
            try:
                html_content = file_content.decode("utf-8", errors="ignore")
                # 简单去除HTML标签
                import re

                text = re.sub(r"<[^>]+>", " ", html_content)
                text = re.sub(r"\s+", " ", text).strip()
                logger.info(f"从HTML简单提取了 {len(text)} 字符的文本")
                return text
            except Exception as e2:
                logger.error(f"HTML简单提取也失败: {e2}")
                return ""

    def extract_text_from_excel_xlrd(self, file_content: bytes) -> str:
        """使用xlrd从老式.xls文件提取文本"""
        if not XLRD_AVAILABLE:
            logger.error("xlrd not available for .xls file processing")
            return ""

        try:
            workbook = xlrd.open_workbook(file_contents=file_content)
            full_text = []

            logger.info(f"📊 .xls工作簿包含 {workbook.nsheets} 个工作表")

            for sheet_index in range(workbook.nsheets):
                sheet = workbook.sheet_by_index(sheet_index)
                sheet_name = sheet.name
                logger.info(f"📄 正在处理工作表: {sheet_name}")

                sheet_text = []
                for row_idx in range(sheet.nrows):
                    row_text = []
                    for col_idx in range(sheet.ncols):
                        cell = sheet.cell(row_idx, col_idx)
                        if cell.value:
                            row_text.append(str(cell.value).strip())
                    if row_text:
                        sheet_text.append(" | ".join(row_text))

                if sheet_text:
                    full_text.append(f"=== {sheet_name} ===")
                    full_text.extend(sheet_text)
                    logger.info(
                        f"✅ 从工作表 {sheet_name} 提取了 {len(sheet_text)} 行数据"
                    )

            text = "\n".join(full_text)
            logger.info(f"✅ 从.xls文档提取了 {len(text)} 字符的文本")
            return text

        except Exception as e:
            logger.error(f"❌ xlrd解析.xls文件失败: {e}")
            return ""

    def extract_text_from_excel_pandas(self, file_content: bytes, filename: str) -> str:
        """使用pandas作为备用方案提取Excel文件"""
        if not PANDAS_AVAILABLE:
            logger.error("pandas not available for Excel processing")
            return ""

        try:
            logger.info(f"🐼 尝试使用pandas解析: {filename}")

            # 根据文件扩展名选择引擎
            if filename.lower().endswith(".xls"):
                # 对于.xls文件使用xlrd引擎
                dfs = pd.read_excel(
                    io.BytesIO(file_content), sheet_name=None, engine="xlrd"
                )
            else:
                # 对于.xlsx文件使用openpyxl引擎
                dfs = pd.read_excel(
                    io.BytesIO(file_content), sheet_name=None, engine="openpyxl"
                )

            full_text = []

            for sheet_name, df in dfs.items():
                logger.info(f"📄 处理工作表: {sheet_name}")
                sheet_text = [f"=== {sheet_name} ==="]

                # 转换DataFrame为文本
                for _, row in df.iterrows():
                    row_text = []
                    for value in row:
                        if pd.notna(value) and str(value).strip():
                            row_text.append(str(value).strip())
                    if row_text:
                        sheet_text.append(" | ".join(row_text))

                full_text.extend(sheet_text)
                logger.info(
                    f"✅ 从工作表 {sheet_name} 提取了 {len(sheet_text)-1} 行数据"
                )

            text = "\n".join(full_text)
            logger.info(f"✅ pandas成功提取了 {len(text)} 字符的文本")
            return text

        except Exception as e:
            logger.error(f"❌ pandas解析Excel文件失败: {e}")
            return ""

    def extract_text_from_excel_openpyxl(self, file_content: bytes) -> str:
        """使用openpyxl从.xlsx文件提取文本"""
        if not OPENPYXL_AVAILABLE:
            logger.error("❌ openpyxl not available for Excel document processing")
            return ""

        try:
            logger.info("🔧 使用openpyxl开始解析Excel文件...")
            workbook = openpyxl.load_workbook(io.BytesIO(file_content))
            full_text = []

            logger.info(
                f"📊 Excel工作簿包含 {len(workbook.sheetnames)} 个工作表: {workbook.sheetnames}"
            )

            for sheet_name in workbook.sheetnames:
                logger.info(f"📄 正在处理工作表: {sheet_name}")
                sheet = workbook[sheet_name]
                sheet_text = []

                for row_num, row in enumerate(sheet.iter_rows(values_only=True), 1):
                    row_text = []
                    for cell in row:
                        if cell is not None and str(cell).strip():
                            row_text.append(str(cell).strip())
                    if row_text:
                        sheet_text.append(" | ".join(row_text))

                if sheet_text:
                    full_text.append(f"=== {sheet_name} ===")
                    full_text.extend(sheet_text)
                    logger.info(
                        f"✅ 从工作表 {sheet_name} 提取了 {len(sheet_text)} 行数据"
                    )

            text = "\n".join(full_text)
            logger.info(f"✅ openpyxl成功提取了 {len(text)} 字符的文本")
            return text

        except Exception as e:
            logger.error(f"❌ openpyxl解析失败: {e}")
            return ""

    def extract_text_from_excel(self, file_content: bytes, filename: str) -> str:
        """智能Excel文件文本提取（支持多种方法和格式）"""
        logger.info(f"🔧 开始智能解析Excel文件: {filename}")

        # 检测文件真实类型
        file_type = self.detect_file_type(file_content, filename)
        logger.info(f"📊 检测到文件类型: {file_type}")

        # 如果检测为HTML表格，直接处理
        if file_type == "html_table":
            logger.info("🌐 检测为HTML表格，使用HTML解析器")
            return self.extract_text_from_html_table(file_content)

        # 根据文件名扩展名和检测结果选择处理方法
        is_xls = filename.lower().endswith(".xls") or file_type == "xls"
        is_xlsx = filename.lower().endswith(".xlsx") or file_type == "xlsx"

        extraction_methods = []

        if is_xlsx:
            # .xlsx文件优先使用openpyxl
            extraction_methods = [
                ("openpyxl", self.extract_text_from_excel_openpyxl),
                (
                    "pandas(.xlsx)",
                    lambda content: self.extract_text_from_excel_pandas(
                        content, filename
                    ),
                ),
            ]
        elif is_xls:
            # .xls文件优先使用xlrd
            extraction_methods = [
                ("xlrd", self.extract_text_from_excel_xlrd),
                (
                    "pandas(.xls)",
                    lambda content: self.extract_text_from_excel_pandas(
                        content, filename
                    ),
                ),
                ("html_table", self.extract_text_from_html_table),  # 一些.xls实际是HTML
            ]
        else:
            # 未知格式，尝试所有方法
            extraction_methods = [
                ("openpyxl", self.extract_text_from_excel_openpyxl),
                ("xlrd", self.extract_text_from_excel_xlrd),
                (
                    "pandas",
                    lambda content: self.extract_text_from_excel_pandas(
                        content, filename
                    ),
                ),
                ("html_table", self.extract_text_from_html_table),
            ]

        # 尝试各种提取方法
        for method_name, method_func in extraction_methods:
            try:
                logger.info(f"🔄 尝试使用 {method_name} 解析...")

                if method_name == "openpyxl" and not OPENPYXL_AVAILABLE:
                    logger.info(f"⏭️ {method_name} 不可用，跳过")
                    continue
                elif method_name == "xlrd" and not XLRD_AVAILABLE:
                    logger.info(f"⏭️ {method_name} 不可用，跳过")
                    continue
                elif method_name.startswith("pandas") and not PANDAS_AVAILABLE:
                    logger.info(f"⏭️ {method_name} 不可用，跳过")
                    continue

                text = method_func(file_content)

                if text and text.strip():
                    logger.info(f"✅ {method_name} 成功提取了 {len(text)} 字符的文本")

                    # 🔧 确保控制台输出
                    print(f"\n{'='*60}")
                    print(f"📊 Excel文件解析结果 (使用 {method_name}):")
                    print(f"文件名: {filename}")
                    print(f"{'='*60}")
                    print(text[:3000] + ("..." if len(text) > 3000 else ""))
                    print(f"{'='*60}\n")

                    return text
                else:
                    logger.warning(f"⚠️ {method_name} 未能提取到有效文本")

            except Exception as e:
                logger.warning(f"❌ {method_name} 解析失败: {e}")
                continue

        # 所有方法都失败了
        logger.error(f"💥 所有Excel解析方法都失败了: {filename}")

        # 最后尝试：看看是否是纯文本文件被错误命名
        try:
            text_content = file_content.decode("utf-8", errors="ignore")
            if len(text_content.strip()) > 50:  # 如果有足够的文本内容
                logger.info("🔤 作为纯文本文件处理")
                return text_content
        except:
            pass

        return ""

    def extract_text_from_attachment(self, attachment: Dict) -> str:
        """根据附件类型提取文本内容 - 改进版"""
        filename = attachment.get("filename", "").lower()
        original_filename = attachment.get("original_filename", "")
        file_content = attachment.get("content", b"")

        logger.info(f"🔧 准备从附件提取文本: {filename}")
        logger.info(f"   原始文件名: {original_filename}")
        logger.info(f"   文件大小: {len(file_content)} 字节")

        if not file_content:
            logger.warning(f"⚠️ 附件 {filename} 没有内容")
            return ""

        # 如果内容是base64编码的字符串，先解码
        if isinstance(file_content, str):
            try:
                file_content = base64.b64decode(file_content)
                logger.info(f"✅ Base64解码成功")
            except Exception as e:
                logger.error(f"❌ Base64解码失败: {e}")
                return ""

        # 检测文件真实类型
        detected_type = self.detect_file_type(file_content, filename)
        logger.info(f"🔍 检测到的文件类型: {detected_type}")

        # 根据文件扩展名和检测结果选择解析方法
        if filename.endswith((".docx", ".doc")) or detected_type == "docx":
            logger.info(f"📄 作为Word文档处理")
            return self.extract_text_from_docx(file_content)
        elif filename.endswith(".pdf") or detected_type == "pdf":
            logger.info(f"📄 作为PDF文档处理")
            return self.extract_text_from_pdf(file_content)
        elif filename.endswith((".xlsx", ".xls")) or detected_type in [
            "xlsx",
            "xls",
            "html_table",
        ]:
            logger.info(f"📊 作为Excel文档处理")
            return self.extract_text_from_excel(file_content, filename)
        elif detected_type == "text":
            logger.info(f"🔤 作为纯文本处理")
            try:
                return file_content.decode("utf-8", errors="ignore")
            except:
                return ""
        else:
            logger.warning(
                f"❓ 不支持的文件类型: {filename} (检测类型: {detected_type})"
            )
            return ""

    def _extract_json_from_text(self, text: str) -> Optional[Dict]:
        """从AI响应文本中提取JSON"""
        try:
            # 尝试直接解析
            result = json.loads(text.strip())
            return result
        except json.JSONDecodeError:
            # 尝试查找JSON块
            json_pattern = r"\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}"
            matches = re.findall(json_pattern, text, re.DOTALL)

            for match in matches:
                try:
                    result = json.loads(match)
                    return result
                except json.JSONDecodeError:
                    continue

            # 尝试复杂的JSON提取
            start_idx = text.find("{")
            if start_idx != -1:
                brace_count = 0
                for i, char in enumerate(text[start_idx:], start_idx):
                    if char == "{":
                        brace_count += 1
                    elif char == "}":
                        brace_count -= 1
                        if brace_count == 0:
                            try:
                                extracted = text[start_idx : i + 1]
                                result = json.loads(extracted)
                                return result
                            except json.JSONDecodeError:
                                break

            logger.warning(f"无法从文本中提取JSON: {text[:200]}...")
            return None

    async def extract_resume_data_with_ai(
        self, resume_text: str, filename: str = ""
    ) -> Optional[ResumeData]:
        """使用AI从简历文本中提取结构化数据"""
        if not self.ai_client:
            logger.warning(
                "AI client not initialized. Skipping resume data extraction."
            )
            return None

        if not resume_text.strip():
            logger.warning("简历文本为空，跳过提取")
            return None

        provider_name = self.ai_config.get("provider_name")
        model_extract = self.ai_config.get("model_extract", "gpt-4")
        temperature = self.ai_config.get("temperature", 0.3)
        max_tokens_extract = self.ai_config.get("max_tokens", 2048)

        # 限制文本长度，避免超出AI模型限制
        if len(resume_text) > 4000:
            resume_text = resume_text[:4000] + "..."

        prompt = f"""
以下は履歴書・職務経歴書の内容です。この情報から技術者の詳細情報を抽出して、必ずJSON形式で返してください。

【ファイル名】: {filename}
【履歴書内容】:
{resume_text}

以下の形式で抽出してください：
{{
    "name": "技術者名（必須）",
    "email": "メールアドレス",
    "phone": "電話番号",
    "gender": "性別（男性/女性/回答しない）",
    "age": "年齢",
    "nationality": "国籍",
    "nearest_station": "最寄り駅",
    "education": "学歴・最終学歴",
    "arrival_year_japan": "来日年度",
    "certifications": ["資格1", "資格2"],
    "skills": ["プログラミング言語1", "プログラミング言語2", "技術スキル1"],
    "technical_keywords": ["Java", "Python", "AWS", "React"],
    "experience": "総経験年数（例：5年）",
    "work_scope": "作業範囲・得意分野",
    "work_experience": "職務経歴の詳細",
    "japanese_level": "日本語レベル（不問/日常会話レベル/ビジネスレベル/ネイティブレベル）",
    "english_level": "英語レベル（不問/日常会話レベル/ビジネスレベル/ネイティブレベル）",
    "availability": "稼働可能時期",
    "preferred_work_style": ["常駐", "リモート", "ハイブリッド"],
    "preferred_locations": ["東京", "大阪"],
    "desired_rate_min": 希望単価下限（数値のみ、万円単位）,
    "desired_rate_max": 希望単価上限（数値のみ、万円単位）,
    "overtime_available": 残業対応可能（true/false）,
    "business_trip_available": 出張対応可能（true/false）,
    "self_promotion": "自己PR・アピールポイント",
    "remarks": "備考・その他",
    "recommendation": "推薦コメント",
    "source_filename": "{filename}"
}}

重要な指示：
1. nameフィールドは必須です。見つからない場合は"名前不明"としてください
2. 情報が見つからない項目はnullにしてください
3. desired_rate_min/maxは数値のみを返してください（万円表記は除く）
4. skillsには具体的な技術名を含めてください
5. technical_keywordsには技術関連のキーワードを抽出してください
6. JSONのみを返してください、他の説明は不要です
"""

        messages = [
            {
                "role": "system",
                "content": "あなたは履歴書・職務経歴書から情報を抽出する専門家です。必ずJSONのみを返してください。",
            },
            {"role": "user", "content": prompt},
        ]

        try:
            if provider_name == "openai":
                response = await self.ai_client.chat.completions.create(
                    model=model_extract,
                    messages=messages,
                    temperature=temperature,
                    max_tokens=max_tokens_extract,
                )
                raw_content = response.choices[0].message.content
                data = self._extract_json_from_text(raw_content)

            elif provider_name in ["deepseek", "custom"]:
                if isinstance(self.ai_client, httpx.AsyncClient):
                    try:
                        logger.info(
                            f"发送简历解析请求到{provider_name.title()} API: {filename}"
                        )

                        response = await self.ai_client.post(
                            "/v1/chat/completions",
                            json={
                                "model": model_extract,
                                "messages": messages,
                                "temperature": temperature,
                                "max_tokens": max_tokens_extract,
                            },
                        )
                        response.raise_for_status()

                        response_json = response.json()
                        raw_response_content = response_json["choices"][0]["message"][
                            "content"
                        ]

                        logger.info(
                            f"=== {provider_name.title()} 简历解析响应 ({filename}) ==="
                        )
                        logger.info(f"Raw content:\n{raw_response_content}")

                        data = self._extract_json_from_text(raw_response_content)
                        if data:
                            logger.info(f"成功解析简历JSON: {filename}")
                        else:
                            logger.error(f"JSON解析失败: {filename}")

                    except Exception as e:
                        logger.error(
                            f"{provider_name.title()} API error for resume {filename}: {e}"
                        )
                        return None
                else:
                    logger.warning(
                        f"{provider_name.title()} client not available for resume extraction"
                    )
                    return None
            else:
                logger.warning(
                    f"Unsupported AI provider for resume extraction: {provider_name}"
                )
                return None

            if data:
                # 确保name字段存在
                if not data.get("name"):
                    data["name"] = "名前不明"

                # 确保experience字段存在
                if not data.get("experience"):
                    data["experience"] = "不明"

                return ResumeData(**data)
            else:
                logger.error(
                    f"Failed to parse JSON from AI response for resume: {filename}"
                )
                return None

        except Exception as e:
            logger.error(f"Error extracting resume data from {filename}: {e}")
            import traceback

            logger.error(f"Full traceback: {traceback.format_exc()}")
            return None

    async def process_resume_attachments(
        self, attachments: List[Dict]
    ) -> List[ResumeData]:
        """处理所有简历附件，返回提取的简历数据列表"""
        resume_data_list = []

        # 过滤出可能的简历文件
        resume_files = []
        resume_extensions = [".docx", ".doc", ".pdf", ".xlsx", ".xls"]
        engineer_patterns = [
            r"履歴書",
            r"職務経歴",
            r"スキルシート",
            r"resume",
            r"cv",
            r"profile",
        ]

        logger.info(f"🔍 开始分析 {len(attachments)} 个附件")

        for attachment in attachments:
            filename = attachment.get("filename", "").lower()
            original_filename = attachment.get("original_filename", "")

            logger.info(f"📄 分析文件: '{filename}' (原始: '{original_filename}')")

            # 检查文件扩展名
            has_resume_extension = any(
                filename.endswith(ext) for ext in resume_extensions
            )

            # 检查文件名关键词
            has_resume_keyword = any(
                re.search(pattern, filename) for pattern in engineer_patterns
            )

            logger.info(f"   扩展名匹配: {has_resume_extension}")
            logger.info(f"   关键词匹配: {has_resume_keyword}")

            if has_resume_extension or has_resume_keyword:
                resume_files.append(attachment)
                logger.info(f"✅ 确认为简历文件: {filename}")

        if not resume_files:
            logger.info("📭 未发现简历附件")
            return resume_data_list

        logger.info(f"📋 开始处理 {len(resume_files)} 个简历文件")

        for attachment in resume_files:
            filename = attachment.get("filename", "")
            logger.info(f"🔄 正在处理简历文件: {filename}")

            try:
                # 提取文本内容
                resume_text = self.extract_text_from_attachment(attachment)

                if not resume_text.strip():
                    logger.warning(f"⚠️ 无法从文件 {filename} 中提取文本内容")
                    continue

                logger.info(f"📝 从 {filename} 提取了 {len(resume_text)} 字符的文本")

                # 使用AI提取结构化数据
                resume_data = await self.extract_resume_data_with_ai(
                    resume_text, filename
                )

                if resume_data:
                    resume_data_list.append(resume_data)
                    logger.info(f"✅ 成功提取简历数据: {resume_data.name} ({filename})")
                else:
                    logger.error(f"❌ 无法从 {filename} 提取简历数据")

            except Exception as e:
                logger.error(f"💥 处理简历文件 {filename} 时出错: {e}")
                import traceback

                logger.error(f"详细错误: {traceback.format_exc()}")
                continue

        logger.info(f"🎯 简历处理完成，成功提取 {len(resume_data_list)} 份简历数据")
        return resume_data_list

    def has_resume_attachments(self, attachments: List[Dict]) -> bool:
        """检查是否包含简历附件"""
        if not attachments:
            logger.info("📎 没有附件")
            return False

        resume_extensions = [".docx", ".doc", ".pdf", ".xlsx", ".xls"]
        engineer_patterns = [
            r"履歴書",
            r"職務経歴",
            r"スキルシート",
            r"resume",
            r"cv",
            r"profile",
        ]

        logger.info(f"📎 检查 {len(attachments)} 个附件是否为简历文件")

        for i, attachment in enumerate(attachments, 1):
            filename = attachment.get("filename", "").lower()
            original_filename = attachment.get("original_filename", "")

            logger.info(f"📄 附件 {i}: '{filename}' (原始: '{original_filename}')")

            # 检查文件扩展名
            has_resume_extension = any(
                filename.endswith(ext) for ext in resume_extensions
            )

            if has_resume_extension:
                logger.info(f"✅ 附件 {i} 匹配简历扩展名")
            else:
                logger.info(f"❌ 附件 {i} 不匹配简历扩展名 {resume_extensions}")

            # 检查文件名关键词
            has_resume_keyword = any(
                re.search(pattern, filename) for pattern in engineer_patterns
            )

            if has_resume_keyword:
                logger.info(f"✅ 附件 {i} 匹配简历关键词")
            else:
                logger.info(f"❌ 附件 {i} 不匹配简历关键词 {engineer_patterns}")

            if has_resume_extension or has_resume_keyword:
                logger.info(f"🎯 确认附件 {i} 为简历文件")
                return True

        logger.info("📭 未发现简历附件")
        return False
